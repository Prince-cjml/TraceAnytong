from __future__ import annotations

import io

from docx import Document
from docx.shared import Inches

from ..carriers.screen_tile import ScreenTileCarrier
from ..errors import InvalidArtifactError
from ..fingerprint.perceptual import PerceptualFingerprinter
from ..models import Artifact, CarrierEvidence, CarrierProfile, PersonalizationResult, TraceIdentity
from .base import BaseAdapter


class DocxFormatAdapter(BaseAdapter):
    mime_types = frozenset({"application/vnd.openxmlformats-officedocument.wordprocessingml.document"})

    def __init__(self, carrier: ScreenTileCarrier | None = None) -> None:
        self.carrier = carrier or ScreenTileCarrier()

    def personalize(self, artifact: Artifact, trace_identity: TraceIdentity, profile: CarrierProfile, **_: object) -> PersonalizationResult:
        self.require_support(artifact.mime_type)
        try:
            document = Document(io.BytesIO(artifact.data))
            tile_file = io.BytesIO()
            self.carrier.tile_rgba(trace_identity, profile).save(tile_file, "PNG")
            # Headers are a durable native render layer and leave body XML editable.
            for section in document.sections:
                header = section.header
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                paragraph.add_run().add_picture(io.BytesIO(tile_file.getvalue()), width=Inches(6.5))
            document.core_properties.comments = f"TraceAnytong opaque trace {trace_identity.trace_handle}; profile {profile.profile_version}"
            out = io.BytesIO()
            document.save(out)
        except Exception as exc:
            raise InvalidArtifactError("DOCX bytes cannot be personalized") from exc
        result = Artifact(out.getvalue(), artifact.mime_type, artifact.filename)
        return PersonalizationResult(result, CarrierEvidence("screen", self.carrier.detector_version, 1.0, {"sections": len(document.sections), "placement": "header-image", "carrierVersion": profile.carrier_version, "provenance": "core-properties-comments"}), PerceptualFingerprinter().index(result))

    def render_reference(self, artifact: Artifact) -> list[Artifact]:
        self.require_support(artifact.mime_type)
        # Rendering DOCX requires a platform office engine; preserving this typed fact is safer than pretending success.
        return []
